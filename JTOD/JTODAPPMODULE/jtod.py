import json
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH


def convert_json_to_docx(filename):
    # Open and Read json file
    if filename.startswith("file://"):
        filename = filename[7:]
    with open(filename) as j:
            data = json.load(j)


    # Create a new document
    doc = Document()

    # Add header text
    header_text = 'Nollydic entry record for ' + filename[-15:-5]
    header_paragraph = doc.add_heading(header_text, 0)
    header_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header_paragraph.add_run().bold = True
    #header = 'Nollydic entry record for ' + filename[-15:-5]
    #header = doc.add_heading(header , 0)
    #header.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add a table
    table = doc.add_table(rows=1, cols=7)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Set the table headers
    headers = ['S/N', 'Name', 'Actual Time', 'Input Time', 'Role', 'Role Info', 'Phone Number']
    hdr_cells = table.rows[0].cells
    for i in range(len(headers)):
        hdr_paragraph = hdr_cells[i].paragraphs[0]
        hdr_paragraph.add_run(headers[i]).bold = True
        #hdr_cells[i].text = headers[i]
        #hdr_cells[i].text.bold = True

    # Add data to the table
    count = 1
    for person_id in data:
        person = data[person_id]
        row_cells = table.add_row().cells
        row_cells[0].text = str(count)
        row_cells[1].text = person['name']
        row_cells[2].text = person['actual_dtime']
        row_cells[3].text = person['input_dtime']
        row_cells[4].text = person['role']
        row_cells[5].text = person['role_info']
        row_cells[6].text = person['phone']
        count += 1



    # Save the Document
    docxname = filename[:-5]
    docxname = docxname + ".docx"
    doc.save(docxname)
