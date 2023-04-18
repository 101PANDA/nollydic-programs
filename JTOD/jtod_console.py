import json
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Collect file name
filename = input("Enter file name: ")

# Check if file is json
if fileext[-5:] != ".json":
    print(filename, "not allowed, file must have a .json extenstion")

# Open and Read json file
with open(filename) as j:
    data = json.load(j)


# Create a new document
doc = Document()

# Add header text
header = 'Nollydic entry record for ' + filename[:-5]
header = doc.add_heading(header , 0)
header.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Add a table
table = doc.add_table(rows=1, cols=7)
table.alignment = WD_TABLE_ALIGNMENT.CENTER

# Set the table headers
headers = ['S/N', 'Name', 'Actual Time', 'Input Time', 'Role', 'Role Info', 'Phone Number']
hdr_cells = table.rows[0].cells
for i in range(len(headers)):
    hdr_cells[i].text = headers[i]

# Add data to the table
count = 1
for person_id in data:
    person = data[person_id]
    row_cells = table.add_row().cells
    row_cells[0].text = str(count)
    row_cells[1].text = person['name']
    row_cells[2].text = person['actual_dtime']
    row_cells[3].text = person['input_dtime']
    row_cells[4].text = person['role']
    row_cells[5].text = person['role_info']
    row_cells[6].text = person['phone']
    count += 1



# Save the Document
docxname = filename[:-5] + ".docx"
doc.save(docxname)
